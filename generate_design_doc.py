"""
生成「嗨 Hai」系统设计文档（Word 格式）。
运行（任选其一）：
  py generate_design_doc.py          # Windows 推荐（Python Launcher）
  python generate_design_doc.py    # 若已配置 PATH
  .\\generate_design_doc.bat       # Windows：自动尝试 py / python
输出：Design/嗨Hai健康管理系统设计文档.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT_DIR = Path(__file__).resolve().parent / "Design"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "嗨Hai健康管理系统设计文档.docx"


def set_cell_shading(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    tcPr.append(shading)


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B6B4A")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    return table


def build_document() -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "SimHei"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        hs.font.color.rgb = RGBColor(0x1B, 0x6B, 0x4A)

    # ===== 封面 =====
    for _ in range(4):
        doc.add_paragraph("")
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("基于多模态大模型的个性化\n全周期健康管理与智能预警系统")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x6B, 0x4A)
    run.font.name = "SimHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("—— 「嗨 Hai」系统设计文档 ——")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x5C, 0x6F, 0x66)
    sr.font.name = "SimHei"
    sr.element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    for _ in range(3):
        doc.add_paragraph("")

    info_lines = [
        "课程名称：医学人工智能",
        "项目名称：嗨 Hai — 全周期健康管理与智能预警平台",
        "日　　期：2026 年 4 月",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)

    doc.add_page_break()

    # ===== 目录占位 =====
    doc.add_heading("目录", level=1)
    doc.add_paragraph("（生成后请在 Word 中插入自动目录：引用 → 目录 → 自动目录）")
    doc.add_page_break()

    # ===== 1. 选题背景与需求分析 =====
    doc.add_heading("1  选题背景与需求分析", level=1)

    doc.add_heading("1.1  研究背景", level=2)
    doc.add_paragraph(
        "随着慢性非传染性疾病（NCD）负担日益加重，世界卫生组织（WHO）数据表明，全球约 74% 的死亡归因于心血管疾病、"
        "糖尿病、慢性呼吸系统疾病与癌症等慢性病 [1]。中国作为慢病高发国家，2 型糖尿病患病率已超过 11%，"
        "高血压患病率约 27.5% [2]。传统的就医模式以「事后治疗」为核心，患者往往在出现显著症状后才寻求医疗干预，"
        "导致疾病管理的最佳窗口期被错过。"
    )
    doc.add_paragraph(
        "近年来，以大语言模型（LLM）为代表的生成式人工智能技术取得突破性进展。GPT-4、DeepSeek、"
        "LLaMA 等模型在自然语言理解与生成任务上展现出接近甚至超越人类的表现 [3]。将这类技术应用于健康管理领域，"
        "有望实现从「被动治疗」到「主动预防与全周期管理」的范式转变。检索增强生成（RAG）技术可在不重新训练模型的前提下，"
        "注入领域特定的知识片段，有效降低大模型的「幻觉」风险 [4]，使其输出更贴合医学科普的严谨性要求。"
    )

    doc.add_heading("1.2  需求分析", level=2)
    doc.add_paragraph(
        "通过调研面向普通用户（非临床专业人员）的健康管理需求，本项目归纳出以下核心需求："
    )
    needs = [
        ("个人健康档案数字化", "用户可录入身高、体重、年龄、饮食与运动偏好、既往病史，并上传体检报告（PDF/TXT），系统自动解析关键检验指标。"),
        ("综合健康评估与可视化", "在总览页以 BMI、综合健康指数、饮食营养雷达图、运动负荷柱图、康复阶段路径等多维度呈现健康状态，便于用户直观理解。"),
        ("智能明日规划", "基于用户档案与昨日回顾（步数、运动、饮食、睡眠等），结合营养学公式与大模型推理，生成次日饮食、运动与作息建议。"),
        ("体征监测与规则预警", "导入或模拟心率、血氧、血压、血糖等时序数据，通过规则引擎检测连续越界趋势，触发可视化预警。"),
        ("AI 科普问答", "多轮对话结合 RAG 知识库，回答用户健康相关疑问，严格遵守「不诊断、不开药」的安全约束。"),
        ("数据安全与审计", "所有数据本地存储（SQLite），操作全程审计日志记录，输入数据经过严格校验，保障数据合规性。"),
    ]
    for title, desc in needs:
        doc.add_paragraph(f"{title}：{desc}", style="List Bullet")

    doc.add_heading("1.3  目标用户与使用场景", level=2)
    doc.add_paragraph(
        "目标用户为具有自我健康管理意识的普通成年人群，包括但不限于：慢病（高血压、糖尿病、血脂异常）患者的日常自我管理、"
        "术后康复期的渐进性活动追踪、以及健康人群的膳食与运动科学规划。系统定位为科普级健康助手，明确声明不替代执业医师面诊。"
    )

    # ===== 2. 相关研究 =====
    doc.add_heading("2  相关研究", level=1)

    doc.add_heading("2.1  大语言模型在医疗健康领域的应用", level=2)
    doc.add_paragraph(
        "大语言模型（LLM）在医疗领域的应用研究近年来迅速增长。Google 的 Med-PaLM 2 在 USMLE 风格的医学问答中达到专家级水平 [5]，"
        "表明 LLM 具备处理复杂医学知识的潜力。然而，直接将通用 LLM 用于医疗建议面临「幻觉」风险——模型可能生成看似合理但事实错误的内容。"
        "为此，RAG（Retrieval-Augmented Generation）架构被广泛采用：先从可信知识库中检索相关片段，再将其注入模型上下文 [4]，"
        "从而在保持生成能力的同时提升事实准确性。"
    )

    doc.add_heading("2.2  个性化健康管理系统", level=2)
    doc.add_paragraph(
        "现有健康管理应用可分为两类：一类以硬件为核心（如 Apple Health、华为健康），依赖可穿戴设备采集生理数据；"
        "另一类以软件为核心（如 MyFitnessPal、薄荷健康），侧重饮食与运动记录。两者均存在局限：前者对非智能设备用户不友好，"
        "后者缺乏 AI 驱动的个性化推理能力。本项目尝试融合两者优势，通过模拟体征数据接入与 LLM 驱动的推理引擎，"
        "构建一个不依赖特定硬件的全周期管理方案。"
    )

    doc.add_heading("2.3  体征预警规则引擎", level=2)
    doc.add_paragraph(
        "临床危急值（Critical Values）管理是医疗安全的核心环节。传统做法基于阈值规则（如心率 >140 或 <45 持续 N 次触发告警）[6]。"
        "本项目借鉴该思路，在本地实现轻量级规则引擎，对心率、血氧、血压与血糖等指标进行连续越界检测。"
        "告警的最终判定不依赖大模型（避免 LLM 误判危重），而由规则引擎确定性输出，大模型仅负责自然语言解释。"
    )

    doc.add_heading("2.4  营养学估算模型", level=2)
    doc.add_paragraph(
        "Mifflin-St Jeor 公式是目前公认较为准确的基础代谢率（BMR）估算方法 [7]。"
        "本项目采用该公式结合活动因子与步数调整系数，估算每日总能量消耗（TDEE）及三大营养素参考摄入量，"
        "作为大模型生成明日饮食建议的量化锚点。"
    )

    # ===== 3. 系统设计 =====
    doc.add_heading("3  系统设计", level=1)

    doc.add_heading("3.1  系统架构概述", level=2)
    doc.add_paragraph(
        "「嗨 Hai」采用「前置流处理 + 后置多模态推理」的分层架构。系统由输入层、预处理层、上下文融合层、"
        "推理层与输出层五部分组成，数据流向为：多模态输入 → 结构化处理 → 档案与知识融合 → Prompt 组装 → LLM 生成 → 可视化呈现。"
    )

    doc.add_paragraph(
        "系统整体数据流如下（对应 Mermaid 流程图）："
    )
    doc.add_paragraph(
        "输入模态（文本咨询 / PDF体检单 / 体征时序模拟） → "
        "前置处理（文本清洗与意图粗分 / OCR正则结构化 / 规则阈值滑动统计） → "
        "上下文融合（用户档案JSON + 历史摘要 + RAG检索Top-K） → "
        "后置推理（Prompt组装 → DeepSeek大模型生成） → "
        "输出（Web可视化展示 + 趋势图/仪表盘 + 免责声明）"
    )

    doc.add_heading("3.2  技术栈选型", level=2)
    add_styled_table(doc,
        ["层级/组件", "技术选型", "说明"],
        [
            ["前端框架", "Streamlit ≥1.38", "Python 原生 Web 框架，快速构建数据驱动的交互界面"],
            ["编程语言", "Python 3.10+", "生态完善，AI/ML 库丰富"],
            ["关系存储", "SQLite 3", "单文件嵌入式数据库，免安装运维"],
            ["大模型接口", "DeepSeek API（OpenAI 兼容）", "通过 REST 调用，支持 deepseek-chat 模型"],
            ["知识检索", "轻量 RAG（关键词重叠打分）", "无向量模型依赖，Markdown 段落级检索"],
            ["可视化", "Plotly ≥5.18", "交互式图表：雷达图、柱状图、仪表盘、时间线"],
            ["PDF 解析", "pypdf ≥4.0", "纯 Python PDF 文本提取"],
            ["数据处理", "Pandas + NumPy", "时序数据处理与统计分析"],
            ["环境管理", "python-dotenv", "密钥安全管理，避免硬编码"],
        ],
    )

    doc.add_heading("3.3  目录结构", level=2)
    tree = (
        "MedAi/\n"
        "├── app.py                    # Streamlit 主入口\n"
        "├── requirements.txt          # Python 依赖\n"
        "├── .env.example              # 环境变量模板\n"
        "├── .streamlit/config.toml    # Streamlit 主题与端口配置\n"
        "├── core/                     # 核心业务逻辑\n"
        "│   ├── __init__.py\n"
        "│   ├── db.py                 # SQLite 数据访问层\n"
        "│   ├── dashboard.py          # 总览页可视化引擎\n"
        "│   ├── health_score.py       # 综合健康评分引擎\n"
        "│   ├── validators.py         # 数据校验与审计日志\n"
        "│   ├── ingest_report.py      # 体检报告解析（PDF/TXT）\n"
        "│   ├── llm_client.py         # DeepSeek API 统一封装\n"
        "│   ├── rag.py                # 轻量 RAG 知识检索\n"
        "│   ├── recommend.py          # 营养学公式估算（BMR/TDEE）\n"
        "│   ├── recovery_path.py      # 专科康复阶段路径生成\n"
        "│   ├── rules.py              # 体征危急值规则引擎\n"
        "│   ├── ui_styles.py          # 全局 UI 样式系统\n"
        "│   └── vitals.py             # 体征序列生成与加载\n"
        "├── data/\n"
        "│   ├── knowledge/            # RAG 知识库（Markdown）\n"
        "│   └── samples/              # 示例数据文件\n"
        "└── Design/                   # 设计文档与架构图\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(tree)
    run.font.size = Pt(9)
    run.font.name = "Consolas"

    # ===== 4. 模块设计 =====
    doc.add_heading("4  模块设计", level=1)

    doc.add_heading("4.1  智能档案模块（Smart Profiling）", level=2)
    doc.add_paragraph(
        "智能档案模块负责用户健康信息的录入、存储与管理。模块包含以下子功能："
    )
    doc.add_paragraph("基础体征录入：支持身高（cm）、体重（kg）、年龄、性别的表单输入，所有输入经过 validators.py 的范围校验（如身高 50-250 cm、体重 20-300 kg），校验失败时阻止写入并给出明确错误提示。", style="List Bullet")
    doc.add_paragraph("饮食与运动偏好：用户可填写饮食忌口（如少油、不吃牛肉）和运动偏好（如膝关节不好避免跑跳），这些偏好将在明日规划中被引用，避免推荐不适宜的方案。", style="List Bullet")
    doc.add_paragraph("病史记录：支持自由文本录入当前/既往病史，系统通过关键词匹配（如「高血压」「糖尿病」「术后」）自动切换总览页的康复阶段路径轨。", style="List Bullet")
    doc.add_paragraph("体检报告解析：支持 PDF/TXT 格式上传，ingest_report.py 通过正则表达式匹配 9 种常见检验项（ALT、AST、TBIL、CREA、GLU、TC、TG、HDL、LDL），自动提取数值并依据预设参考范围标记「偏高/偏低/正常」。", style="List Bullet")
    doc.add_paragraph("紧急联系人管理：支持添加/删除紧急联系人（姓名+电话），在体征预警触发时展示通知文案（演示环境不发送真实通知）。", style="List Bullet")
    doc.add_paragraph("档案完整度追踪：health_score.py 计算 7 个核心字段的填写率，在总览页以进度条可视化展示，引导用户补全信息。", style="List Bullet")

    doc.add_heading("4.2  综合健康评估模块（Health Scoring）", level=2)
    doc.add_paragraph(
        "本模块是 V2 版本新增的核心升级，实现多维度加权的综合健康指数计算。评分体系如下："
    )
    add_styled_table(doc,
        ["评估维度", "权重", "数据来源", "评分逻辑"],
        [
            ["BMI 体型", "20%", "智能档案（身高/体重）", "18.5-24.0 正常范围得 95 分，超重/肥胖递减"],
            ["运动活力", "25%", "昨日回顾（步数/有氧/力量/睡眠）", "步数、有氧时长、力量时长与睡眠质量综合"],
            ["饮食均衡", "20%", "昨日回顾（正餐/外食/饮食内容）", "在家就餐、规律三餐得分高，外食比例高扣分"],
            ["体检关注", "15%", "检验摘要（异常项）", "有体检记录得 75 分基线，无记录 50 分"],
            ["档案完整", "20%", "智能档案（7 字段填写率）", "直接采用档案完整度百分比"],
        ],
    )
    doc.add_paragraph(
        "评分结果通过 Plotly Gauge 仪表盘（0-100 分）与水平条形图（各维度明细）在总览页顶部展示，"
        "使用红/黄/绿三色分级（<60 红色、60-75 黄色、>75 绿色）直观传达健康状态。"
    )

    doc.add_heading("4.3  明日规划模块（Tomorrow Planner）", level=2)
    doc.add_paragraph(
        "明日规划模块是系统的核心智能功能，实现「昨日回顾 → 公式估算 → LLM 推理 → 明日方案」的闭环。"
    )
    doc.add_paragraph("病情与体征自报：支持体温（℃）、血压（收缩压/舒张压 mmHg）、血糖（mmol/L）的数值录入，所有数值经过 validators.py 的合理范围校验与交叉校验（如收缩压必须高于舒张压）。", style="List Bullet")
    doc.add_paragraph("昨日活动回顾：通过步数（数字输入+滑块同步）、工作活动强度、有氧/无氧时长与类型、正餐顿数、外食比例、睡眠时长、久坐时长等多维度采集。", style="List Bullet")
    doc.add_paragraph("营养学公式估算：recommend.py 采用 Mifflin-St Jeor 公式计算 BMR，乘以活动因子得到 TDEE，结合步数调整系数，输出蛋白质/脂肪/碳水化合物的参考克数。", style="List Bullet")
    doc.add_paragraph("RAG 知识检索：rag.py 从 data/knowledge/ 目录加载 Markdown 知识片段，通过关键词重叠打分检索 Top-5 相关条目，注入大模型上下文。", style="List Bullet")
    doc.add_paragraph("LLM 方案生成：llm_client.py 将个人档案、昨日回顾、公式参考 JSON 与知识库摘录组装为结构化 Prompt，调用 DeepSeek API 生成包含饮食、运动、步数与作息建议的明日方案。", style="List Bullet")

    doc.add_heading("4.4  体征与预警模块（Vitals & Alert）", level=2)
    doc.add_paragraph(
        "本模块实现体征时序数据的可视化与规则引擎预警："
    )
    doc.add_paragraph("数据来源：支持加载示例 CSV、生成平稳随机数据、注入指定类型异常（心动过速/过缓、低血氧、高血糖、高血压）。vitals.py 自动补齐缺失列（血压、血糖），统一为 5 指标时序。", style="List Bullet")
    doc.add_paragraph("可视化：用户可多选指标绘制折线图，数据表展示最近 15 条记录，列名中英文映射。", style="List Bullet")
    doc.add_paragraph("规则引擎（rules.py）：采用滑动窗口（默认 5 点）检测连续越界。阈值设定：心率 >140 或 <45、血氧 <90%、血压 >180/110 mmHg、血糖 >16.7 mmol/L。告警由规则确定性输出，不依赖 LLM 判断。", style="List Bullet")
    doc.add_paragraph("预警响应：触发时显示红色错误横幅与紧急联系人通知文案（演示环境），明确标注「演示规则，非真实急救判定」。", style="List Bullet")

    doc.add_heading("4.5  AI 科普助手模块（Chat Assistant）", level=2)
    doc.add_paragraph(
        "AI 助手模块实现多轮对话式健康科普问答："
    )
    doc.add_paragraph("RAG 增强：每次用户提问前，rag.py 检索 Top-4 知识库片段，以「知识库摘录」形式注入系统提示词。", style="List Bullet")
    doc.add_paragraph("安全约束：System Prompt 硬编码四条约束——不诊断、不开药、优先引用知识库、建议就医。", style="List Bullet")
    doc.add_paragraph("离线兜底：当 DeepSeek API 不可用时，自动切换为规则化离线回复模板，确保系统可用性。", style="List Bullet")
    doc.add_paragraph("历史上下文：最近 10 轮对话注入 messages 列表，支持连续追问。", style="List Bullet")

    doc.add_heading("4.6  数据校验与审计模块（Validators & Audit）", level=2)
    doc.add_paragraph(
        "V2 版本新增的数据治理模块，保障数据录入质量与操作可追溯性："
    )
    doc.add_paragraph("档案校验：身高、体重、年龄、性别的合理范围检查，超出范围阻止写入。", style="List Bullet")
    doc.add_paragraph("体征校验：体温、血压、血糖的范围与交叉逻辑校验（如收缩压 > 舒张压），并对临床警示值（如血糖 ≥16.7）发出额外提醒。", style="List Bullet")
    doc.add_paragraph("审计日志：audit_log 表记录所有关键操作（保存档案、上传报告、同步回顾、AI 方案生成），包含操作名称、模块、详情 JSON 与 UTC 时间戳。", style="List Bullet")
    doc.add_paragraph("日志查阅：总览页底部的折叠区域展示最近 20 条审计记录，便于用户回溯操作历史。", style="List Bullet")

    doc.add_heading("4.7  康复阶段路径模块（Recovery Path）", level=2)
    doc.add_paragraph(
        "recovery_path.py 根据档案病史、检验异常与病情自报，动态生成专科化康复阶段时间线。"
        "系统支持 7 种路径轨：一般健康、高血压管理、糖尿病管理、术后康复、血脂管理、多病共存与慢病综合管理。"
        "每条路径包含 3-5 个阶段节点，每个节点提供标题、摘要与具体建议。当前推荐阶段由规则引擎根据自报数据信号"
        "（发热、高血糖、高血压等）与数据完整度综合推断。可视化采用 Plotly 横向时间线，支持用户点选切换节点。"
    )

    # ===== 5. UI 设计 =====
    doc.add_heading("5  UI 设计", level=1)

    doc.add_heading("5.1  设计理念", level=2)
    doc.add_paragraph(
        "UI 设计遵循「Soft UI + 医疗绿」的 Wellness 风格，参考 UI/UX Pro Max 设计规范。"
        "核心原则：① WCAG 2.1 AA 级对比度（≥4.5:1）；② 150-300ms 微交互动效（hover 抬升 + 阴影加深）；"
        "③ 减少装饰噪音，信息层级清晰；④ 免责声明弱化但不失可读性。"
    )

    doc.add_heading("5.2  色彩体系", level=2)
    add_styled_table(doc,
        ["色彩变量", "色值", "用途"],
        [
            ["--hai-green-800（主色）", "#14523a", "标题、品牌、活跃态按钮"],
            ["--hai-green-700", "#1b6b4a", "主强调色、图表主色"],
            ["--hai-green-500", "#2d8f63", "Focus ring、次级按钮"],
            ["--hai-green-50（背景）", "#f0faf4", "全局页面背景"],
            ["--hai-surface", "#ffffff", "卡片、面板背景"],
            ["--hai-muted", "#5c6f66", "辅助文字、图表标签"],
            ["警告黄", "#c9a227", "中等风险指标"],
            ["危险红", "#c0392b", "高风险指标、错误提示"],
        ],
    )

    doc.add_heading("5.3  组件规范", level=2)
    doc.add_paragraph("圆角体系：卡片 14px，大面板 18px，按钮 10px，徽章 999px（胶囊形）。", style="List Bullet")
    doc.add_paragraph("阴影层级：sm（1px 2px 5%）用于静态卡片，md（4px 20px 7%）用于 hover 态与浮层。", style="List Bullet")
    doc.add_paragraph("动效规范：统一 220ms cubic-bezier(0.25, 0.1, 0.25, 1)，hover 抬升 translateY(-1px ~ -2px)。", style="List Bullet")
    doc.add_paragraph("字体栈：中文优先 Microsoft YaHei / PingFang SC / Noto Sans SC，西文 DM Sans，回退 sans-serif。", style="List Bullet")
    doc.add_paragraph("指标卡片网格：auto-fit 响应式布局（minmax 180px），支持 label/value/sub 三级信息层。", style="List Bullet")

    doc.add_heading("5.4  页面布局", level=2)
    doc.add_paragraph(
        "系统采用侧栏导航 + 主内容区的经典后台布局。侧栏展示品牌标识（嗨 Hai · 全周期健康助手）与 5 个全宽导航按钮"
        "（Material 图标 + 文字）。主内容区宽度限制 1280px，顶部为页面标题与副标题。"
        "总览页布局层级为：指标卡片网格 → 档案完整度进度条 → 健康仪表盘/维度柱图 → 饮食雷达/运动柱图 → 康复阶段时间线 → 审计日志。"
    )

    # ===== 6. 软件接口设计 =====
    doc.add_heading("6  软件接口设计", level=1)

    doc.add_heading("6.1  DeepSeek API 接口", level=2)
    doc.add_paragraph(
        "系统通过 HTTP POST 调用 DeepSeek 的 OpenAI 兼容接口（/v1/chat/completions），"
        "密钥从 .env 文件读取（DEEPSEEK_API_KEY），绝不硬编码于代码库。"
    )
    add_styled_table(doc,
        ["参数", "类型", "说明"],
        [
            ["model", "string", "模型标识，默认 deepseek-chat"],
            ["messages", "array", "对话消息列表（system + user + assistant）"],
            ["temperature", "float", "生成随机性，默认 0.6（兼顾创造力与稳定性）"],
            ["timeout", "int", "请求超时 120 秒"],
        ],
    )

    doc.add_heading("6.2  内部模块接口", level=2)
    add_styled_table(doc,
        ["模块", "核心函数", "输入", "输出"],
        [
            ["db.py", "upsert_profile()", "身高/体重/年龄等字段", "写入 SQLite profile 表"],
            ["db.py", "insert_lab_report()", "文件名/结构化 JSON", "返回记录 ID"],
            ["health_score.py", "compute_health_index()", "档案/昨日回顾/体检/预警", "总分 + 各维度分数"],
            ["validators.py", "validate_profile()", "身高/体重/年龄/性别", "ValidationResult(ok/errors/warnings)"],
            ["validators.py", "validate_vitals_self_report()", "体温/血压/血糖", "ValidationResult"],
            ["validators.py", "write_audit()", "操作名/模块/详情", "写入 audit_log 表"],
            ["recommend.py", "compute_reference_bundle()", "体重/身高/年龄/性别/活动/步数", "BMR/TDEE/宏量营养素 JSON"],
            ["rag.py", "retrieve()", "查询文本 + top_k", "知识片段列表"],
            ["rules.py", "check_critical_vitals()", "体征 DataFrame", "(是否告警, 说明)"],
            ["recovery_path.py", "build_recovery_nodes()", "病史/昨日/检验/是否有报告", "(节点列表, 当前索引, 轨类型)"],
            ["llm_client.py", "chat()", "用户问题/上下文/历史", "(回答文本, 元数据)"],
            ["llm_client.py", "generate_tomorrow_plan()", "档案/回顾/公式/RAG", "(方案 Markdown, 元数据)"],
        ],
    )

    doc.add_heading("6.3  数据库 Schema", level=2)
    add_styled_table(doc,
        ["表名", "主要字段", "用途"],
        [
            ["profile", "id(PK=1), height_cm, weight_kg, age, sex, diet_preferences, medical_history, exercise_preferences, updated_at", "唯一用户档案"],
            ["emergency_contact", "id(PK), name, phone", "紧急联系人"],
            ["lab_report", "id(PK), source_name, raw_text_snippet, structured_json, created_at", "体检报告存档"],
            ["daily_wellness_log", "id(PK), payload_json, created_at", "昨日回顾快照"],
            ["vitals_snapshot", "id(PK), label, payload_json, created_at", "体征序列快照"],
            ["audit_log", "id(PK), action, module, detail_json, created_at", "操作审计日志"],
        ],
    )

    # ===== 7. 可行性分析 =====
    doc.add_heading("7  可行性分析", level=1)

    doc.add_heading("7.1  技术可行性", level=2)
    doc.add_paragraph(
        "系统完全基于 Python 生态构建，核心依赖（Streamlit、Pandas、Plotly、SQLite）均为成熟开源项目，"
        "社区活跃、文档完善。DeepSeek API 提供 OpenAI 兼容接口，切换成本极低。"
        "RAG 采用轻量关键词匹配而非向量模型，消除了 GPU 与大型 Embedding 模型的部署要求。"
        "整套系统可在普通 Windows/macOS/Linux 笔记本上运行，无需 GPU。"
    )

    doc.add_heading("7.2  经济可行性", level=2)
    doc.add_paragraph(
        "系统运行成本极低：所有基础组件免费开源，唯一可选付费项为 DeepSeek API 调用费用。"
        "在课程演示场景下，日均 API 调用次数有限（预估 50-100 次），费用可忽略。"
        "数据存储采用 SQLite 单文件方案，免除数据库服务器租赁费用。系统无需独立服务器部署，"
        "在本地 localhost 运行即可完成演示。"
    )

    doc.add_heading("7.3  操作可行性", level=2)
    doc.add_paragraph(
        "系统面向非技术用户设计，操作流程简洁明了：① 复制 .env.example 并填写 API 密钥；"
        "② 执行 pip install -r requirements.txt 安装依赖；③ streamlit run app.py 启动服务。"
        "前端采用 Streamlit 原生组件（表单、滑块、文件上传、图表等），用户无需学习成本。"
        "数据校验模块在不合理输入时主动阻止写入并给出中文错误提示，降低误操作风险。"
    )

    doc.add_heading("7.4  法律与伦理可行性", level=2)
    doc.add_paragraph(
        "系统定位为课程作业级科普工具，明确声明「不提供诊断、治疗或急救决策」。"
        "每个功能页面底部均展示免责声明。体征预警由规则引擎而非 LLM 判定，避免模型误判危重。"
        "所有数据存储在本地 SQLite，不上传至任何远程服务器。演示数据为合成脱敏数据，不涉及真实患者信息。"
        "API 密钥通过 .env 文件管理并列入 .gitignore，防止泄露。"
    )

    # ===== 8. 效率与性能分析 =====
    doc.add_heading("8  效率与性能分析", level=1)

    doc.add_heading("8.1  响应时间分析", level=2)
    add_styled_table(doc,
        ["操作", "预期响应时间", "瓶颈分析"],
        [
            ["页面加载（总览）", "<1 秒", "本地 SQLite 查询 + Plotly 图表渲染"],
            ["档案保存", "<0.1 秒", "SQLite 单行 UPSERT"],
            ["体检报告解析", "0.2-0.5 秒", "PDF 文本提取 + 正则匹配"],
            ["RAG 检索", "<0.05 秒", "内存关键词匹配，无网络开销"],
            ["AI 方案生成", "3-8 秒", "DeepSeek API 网络延迟 + 生成时间"],
            ["AI 对话回答", "2-5 秒", "同上"],
            ["规则引擎预警", "<0.01 秒", "Pandas 尾部切片 + 布尔运算"],
            ["健康评分计算", "<0.01 秒", "纯算术运算"],
        ],
    )

    doc.add_heading("8.2  存储效率", level=2)
    doc.add_paragraph(
        "SQLite 数据库在典型使用场景下（1 个用户档案、50 条昨日回顾、10 份体检报告、500 条审计日志）"
        "预估占用空间 < 1 MB。body 字段采用 JSON 序列化存储，兼顾灵活性与查询效率。"
        "RAG 知识库为 4 个 Markdown 文件，总计约 1 KB，全量加载至内存后检索延迟可忽略。"
    )

    doc.add_heading("8.3  可扩展性", level=2)
    doc.add_paragraph(
        "若产品化，系统具备清晰的演进路径："
    )
    doc.add_paragraph("数据层：SQLite → PostgreSQL/MySQL，支持多用户并发。", style="List Bullet")
    doc.add_paragraph("检索层：关键词匹配 → FAISS/ChromaDB 向量检索，支持大规模知识库。", style="List Bullet")
    doc.add_paragraph("模型层：DeepSeek API → 本地部署 Ollama/vLLM，降低延迟与费用。", style="List Bullet")
    doc.add_paragraph("流处理：Python 内存队列 → Golang + Kafka，支持高频实时体征接入。", style="List Bullet")
    doc.add_paragraph("前端：Streamlit → React/Vue SPA，支持移动端响应式设计。", style="List Bullet")

    # ===== 参考文献 =====
    doc.add_heading("参考文献", level=1)
    refs = [
        "[1] World Health Organization. Noncommunicable diseases: Key facts. WHO Fact Sheets, 2023.",
        "[2] 中国心血管健康与疾病报告编写组. 中国心血管健康与疾病报告 2022. 中国循环杂志, 2023, 38(6): 583-612.",
        "[3] OpenAI. GPT-4 Technical Report. arXiv preprint arXiv:2303.08774, 2023.",
        "[4] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS, 2020.",
        "[5] Singhal K, Azizi S, Tu T, et al. Large Language Models Encode Clinical Knowledge. Nature, 2023, 620: 172-180.",
        "[6] Lundberg G D. When to panic over abnormal values. Medical Laboratory Observer, 1972, 4: 47-54.",
        "[7] Mifflin M D, St Jeor S T, Hill L A, et al. A new predictive equation for resting energy expenditure in healthy individuals. American Journal of Clinical Nutrition, 1990, 51(2): 241-247.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    # ===== 分工 =====
    doc.add_heading("团队分工", level=1)
    doc.add_paragraph("")
    add_styled_table(doc,
        ["姓名", "学号", "负责模块", "具体工作内容"],
        [
            ["（待填写）", "", "", ""],
            ["（待填写）", "", "", ""],
            ["（待填写）", "", "", ""],
            ["（待填写）", "", "", ""],
            ["（待填写）", "", "", ""],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph("（请各组员填写上表。）")

    return doc


def main():
    import sys
    doc = build_document()
    doc.save(str(OUT_PATH))
    size_kb = OUT_PATH.stat().st_size / 1024
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print(f"Design doc generated: {OUT_PATH}")
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
